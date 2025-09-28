from docx import Document

def extract_content(file_path):
    """
    提取需求文档内容，包括段落、表格、页眉和页脚。
    :param file_path: docx 文件路径
    :return: dict, 包括段落、表格、页眉和页脚
    """
    doc = Document(file_path)
    content = {"paragraphs": [], "tables": [], "headers": [], "footers": []}

    # 提取段落
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content["paragraphs"].append(paragraph.text.strip())

    # 提取表格内容
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)

    # 提取页眉和页脚（如果有）
    if doc.sections:
        for section in doc.sections:
            if section.header:
                content["headers"].append(section.header.paragraphs[0].text.strip())
            if section.footer:
                content["footers"].append(section.footer.paragraphs[0].text.strip())

    return content

import re

def clean_headers_footers(content):
    """
    清理页眉、页脚和目录内容
    :param content: 文档内容字典
    :return: 清理后的段落和表格
    """
    cleaned_content = {"paragraphs": [], "tables": content["tables"]}

    # 清理页眉和页脚中的无用内容
    headers_footers = content["headers"] + content["footers"]
    for paragraph in content["paragraphs"]:
        # 如果段落在页眉或页脚中，过滤掉
        if paragraph in headers_footers:
            continue
        # 过滤掉可能是目录的内容
        if re.match(r"^\s*(第[\d一二三四五六七八九十]+章|\d+(\.\d+)*).*$", paragraph):
            continue
        cleaned_content["paragraphs"].append(paragraph)

    return cleaned_content


template_phrases = ["请填写", "此处为模板", "示例", "模板内容", "请参考以下格式"]

def remove_template_phrases(content):
    """
    删除文档中的模板固有内容
    :param content: 文档段落列表
    :return: 清理后的段落列表
    """
    cleaned_paragraphs = []
    for paragraph in content["paragraphs"]:
        if any(phrase in paragraph for phrase in template_phrases):
            continue  # 跳过模板内容
        cleaned_paragraphs.append(paragraph)
    return {"paragraphs": cleaned_paragraphs, "tables": content["tables"]}


import json

def save_to_json(content, output_path):
    """
    将清理后的文档内容保存为 JSON 格式
    :param content: 清理后的文档内容
    :param output_path: 输出文件路径
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(content, f, ensure_ascii=False, indent=4)


from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QVBoxLayout, QLabel, QTextEdit, QPushButton, QWidget
)

class DocxCleanerGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("智能清理需求文档工具")
        self.setGeometry(200, 200, 800, 600)
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()

        # 文件选择按钮
        self.select_file_btn = QPushButton("选择需求文档")
        self.select_file_btn.clicked.connect(self.select_file)
        layout.addWidget(self.select_file_btn)

        # 文件内容预览
        self.file_content_preview = QTextEdit()
        self.file_content_preview.setReadOnly(True)
        layout.addWidget(QLabel("原始文档内容："))
        layout.addWidget(self.file_content_preview)

        # 清理后内容预览
        self.cleaned_content_preview = QTextEdit()
        self.cleaned_content_preview.setReadOnly(False)
        layout.addWidget(QLabel("清理后文档内容："))
        layout.addWidget(self.cleaned_content_preview)

        # 清理按钮
        self.clean_btn = QPushButton("清理文档")
        self.clean_btn.clicked.connect(self.clean_file)
        layout.addWidget(self.clean_btn)

        # 保存按钮
        self.save_btn = QPushButton("保存清理结果")
        self.save_btn.clicked.connect(self.save_file)
        layout.addWidget(self.save_btn)

        # 设置中心布局
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择DOCX文件", "", "Word文件 (*.docx)")
        if file_path:
            self.file_path = file_path
            content = extract_content(file_path)
            self.file_content_preview.setText("\n".join(content["paragraphs"]))

    def clean_file(self):
        if hasattr(self, "file_path"):
            content = extract_content(self.file_path)
            cleaned_content = clean_headers_footers(content)
            cleaned_content = remove_template_phrases(cleaned_content)
            self.cleaned_content_preview.setText("\n".join(cleaned_content["paragraphs"]))

    def save_file(self):
        if hasattr(self, "file_path"):
            output_path, _ = QFileDialog.getSaveFileName(self, "保存清理结果", "", "JSON文件 (*.json)")
            if output_path:
                content = {"paragraphs": self.cleaned_content_preview.toPlainText().split("\n")}
                save_to_json(content, output_path)

if __name__ == "__main__":
    app = QApplication([])
    window = DocxCleanerGUI()
    window.show()
    app.exec_()
